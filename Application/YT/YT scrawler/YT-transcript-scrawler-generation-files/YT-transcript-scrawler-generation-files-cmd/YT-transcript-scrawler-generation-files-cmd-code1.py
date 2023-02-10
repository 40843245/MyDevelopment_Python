from __future__ import absolute_import

from lxml import etree

import re

import docx

from docx.compat import Unicode
from docx.oxml import OxmlElement
from docx.oxml.exceptions import InvalidXmlError
from docx.oxml.ns import NamespacePrefixedTag, nsmap, qn
from docx.shared import lazyproperty

from docx.oxml import ns

import sys

import json

import docx

from docx import shared

from docx import Document

from docx.shared import Inches

from docx import section

from docx.enum.section import WD_SECTION

from youtube_transcript_api import YouTubeTranscriptApi

from googletrans import Translator

import scrapetube

import webbrowser

import ORDER_class
        
class YT_video():
    
    def Print_Info(self,o,msg):
        print(msg+"=")
        print(o)
        print("type("+msg+")=")
        print(type(o))
        
    def __dict__(self):
        dic={
            "self.transcript_languages=":self.transcript_languages
            ,"self.base_channel_url=":self.base_channel_url
            ,"self.tar_channel_url=":self.tar_channel_url
            ,"self.wantToOpen=":self.wantToOpen
            ,"self.document_name=":self.document_name
            ,"self.document_ext=":self.document_ext
            ,"self.naming_method=":self.naming_method
            ,"self.tar_languanges=":self.tar_languanges
            ,"self.lines=":self.lines
            ,"self.headerList=":self.headerList
            ,"self.footerList=":self.footerList
            ,"self.fontList1=":self.fontList1
            ,"self.fontList2=":self.fontList2
            ,"self.base_watch_video_url=":self.base_watch_video_url
            ,"self.base_channel_url=":self.base_channel_url
            ,"self.videos=":self.videos
            ,"self.video_url_list=":self.video_url_list
        }
        return dic
    def Print_Dict_ElemByElem(self):
        print("The method of the class named Print_Dict_ElemByElem was called.")
        dic=self.__dict__()
        for k,v in dic.items():
            print("key:value,"+str(k)+":"+str(v))
            print("type(key):type(value),"+str(type(k))+":"+str(type(v)))
                    
        
    def __init__(self,transcript_languages,base_channel_url
                 ,tar_channel_url
                 ,wantToOpen,document_name: str
                 ,document_ext : list 
                 ,naming_method : ORDER_class.ORDER
                 ,tar_languanges : list
                 ,lines:int
                 ,headerList:list
                 ,footerList:list
                 ,fontList1:list
                 ,fontList2:list
                 ):
        base_watch_video_url='https://www.youtube.com/watch?v='
        self.base_watch_video_url=base_watch_video_url
        self.transcript_languages = transcript_languages
        self.base_channel_url = base_channel_url
        self.tar_channel_url = tar_channel_url
        self.wantToOpen = wantToOpen
        self.document_name = document_name
        self.document_ext = document_ext
        self.naming_method = naming_method
        self.tar_languanges = tar_languanges
        self.lines=lines
        self.headerList=headerList
        self.footerList=footerList
        self.fontList1=fontList1
        self.fontList2=fontList2
        
        self.videos=list()
        self.video_url_list=list()
        
        self.Set_All_DefaultValues()
        
        
        self.TryOpenUrl(self.base_channel_url,self.tar_channel_url,self.wantToOpen)
        self.GetChannel()
        self.GetAllVideoUrl_List()
        
        print("----------------------")
        print("dict(self)=")
        x=self.__dict__()
        print(x)
        print("----------------------")
         
        
        self.Print_Dict_ElemByElem()
        
        self.YT_handler()
    
    def Set_DefaultValue(self,tar,defaultValue):
        if isinstance(tar,(list,tuple,set)):
            if len(tar)>=1:
                return tar
        
        if isinstance(tar,str):
            if tar!="" and len(tar)==0:
                return tar
        
        if isinstance(tar,(list,tuple)):
            tar=defaultValue
            return tar
        
        if isinstance(tar,(str,int,bool)):
            tar=defaultValue
            return tar
        raise TypeError
        
    def Set_All_DefaultValues(self):       
        self.transcript_languages = self.Set_DefaultValue(self.transcript_languages,['ja'])
        self.base_channel_url = self.Set_DefaultValue(self.base_channel_url,"https://www.youtube.com/channel/")
        self.tar_channel_url = self.Set_DefaultValue(self.tar_channel_url,"UCO-n4ZDDXKPKK29c5eaytpA")
        self.wantToOpen = self.Set_DefaultValue(self.wantToOpen,1)
        self.document_name = self.Set_DefaultValue(self.document_name,"Nanaten")
        self.document_ext = self.Set_DefaultValue(self.document_ext,[".docx"])
        self.naming_method = self.Set_DefaultValue(self.naming_method,ORDER_class.ORDER.ASCENDING_FROM_1)
        self.tar_languanges = self.Set_DefaultValue(self.tar_languanges,['en','zh-tw','ja'])
        
    def ConcatentateListToLetter(self,otherOptions):
        insertedText=""
        if len(otherOptions)<0:
            insertedText=""
            raise "Too less options in the variable otherOptions!!!"
            
        if len(otherOptions)==0:
            insertedText="/"
            return insertedText
        if len(otherOptions)==1:
            insertedText=otherOptions[0]
            return insertedText
        
        for s in otherOptions:
            insertedText=insertedText+str(s)
        return insertedText
    
    def create_element(self,name):
        return OxmlElement(name)

    def create_attribute(self,element, name, value):
        element.set(ns.qn(name), value)
        
    def add_page_number(self,run,options):
        option0=options[0]
        # For other
        if option0!=0 and option0!=1:
            
            otherOptions=options[1:]
            insertedText=self.ConcatentateListToLetter(otherOptions)
            
            fldChar1 = self.create_element('w:fldChar')
            self.create_attribute(fldChar1, 'w:fldCharType', 'begin')
            
            fldChar2 = self.create_element('w:t')
            fldChar2.text=insertedText
            self.create_attribute(fldChar2, 'w:fldCharType', 'separate')

            fldChar3 = self.create_element('w:fldChar')
            self.create_attribute(fldChar2, 'w:fldCharType', 'end')
            
            run._r.append(fldChar1)
            run._r.append(fldChar2)
            run._r.append(fldChar3)
            
        ### For current page and total page
        else:
            fldChar1 = self.create_element('w:fldChar')
            self.create_attribute(fldChar1, 'w:fldCharType', 'begin')
            
            instrText = self.create_element('w:instrText')
            self.create_attribute(instrText, 'xml:space', 'preserve')
        
        
            if option0==0:
                instrText.text = "PAGE"
            elif option0==1:
                instrText.text = "NUMPAGES"
    
            fldChar2 = self.create_element('w:fldChar')
            self.create_attribute(fldChar2, 'w:fldCharType', 'end')

            run._r.append(fldChar1)
            run._r.append(instrText)
            run._r.append(fldChar2)
    def GetChannel(self):
        videos = scrapetube.get_channel(self.tar_channel_url)
        #self.videos=videos
        setattr(self, 'videos', videos)
        
    def GetAllVideoUrl_List(self):
        #self.video_url_list=list()
        v_list=list()
        for video in self.videos:
            v_url=video['videoId']
            video_url=str(self.base_watch_video_url)+v_url
            print("video_url=")
            print(video_url)
            #self.video_url_list.append(video_url)
            v_list.append(video_url)
        print("v_list")
        print(v_list)
        setattr(self, 'video_url_list', v_list)
        
    
    def YT_handler(self):
        print("The method in the class named YT_handler was called.")

        cnt=0

        ## DON'T USE this statement.
        #for video_url_elem in self.videos: 
        ## DON'T directly use the attribute of the class since it is generation expression.
        ## Convert it to the desired type before used.
        
        ## DON'T USE this statement.
        #for video_url_elem in list(self.videos):
        ## Since list(self.videos) is a generation expression and it returns an empty list.
        
        
        ## Such as this statement
        for video_url_elem in self.__dict__()['self.video_url_list=']:
            print("video_url_elem=")
            print(video_url_elem)
            cnt+=1
            
            """
            self.YT_Video_Transcript(video_url=video_url_elem,
                                         document_name=self.document_name+str(cnt),
                                         document_ext=self.document_ext,
                                         transcript_language=self.transcript_languages,
                                         tar_languages=self.tar_languanges,
                                         lines=self.lines
                                         ,headerList=self.headerList
                                         ,footerList=self.footerList
                                         ,fontList1=self.fontList1
                                         ,fontList2=self.fontList2                                                                                                             
                                         )
            """
            
            self.YT_Video_Transcript(video_url=video_url_elem,
                                         document_name=self.__dict__()['self.document_name=']+str(cnt),
                                         document_ext=self.__dict__()['self.document_ext='],
                                         transcript_language=self.__dict__()['self.transcript_languages='],
                                         tar_languages= self.__dict__()['self.tar_languanges='],
                                         lines= self.__dict__()['self.lines=']
                                         ,headerList= self.__dict__()['self.headerList=']
                                         ,footerList= self.__dict__()['self.footerList=']
                                         ,fontList1=self.__dict__()['self.fontList1=']
                                         ,fontList2=self.__dict__()['self.fontList2=']
                                         )
            
            
    def YT_Video_Transcript(self,video_url,
                            document_name
                            ,document_ext
                            ,transcript_language
                            ,tar_languages
                            ,lines:int
                            ,headerList
                            ,footerList
                            ,fontList1
                            ,fontList2
                            ):
        translator = Translator()
        document = Document()
        
        video_id=video_url.split('=')[1]
        print(video_id)
        video_info=YouTubeTranscriptApi.get_transcript(video_id,languages=transcript_language)
        print('-----------------------')
        
        start_index=7
        start_index=start_index+1
        
        msg="The transcript of YT, link="+str(video_url)
        print(msg)
        
        document.add_paragraph(msg)
        document.add_page_break()
        
        self.Print_Info(video_info,"video_info")
        self.Print_Info(lines,"lines")
        
        
        count=0
        for v_info in video_info:
            count+=1
            for la in tar_languages:
                backup=v_info
                msg_para="The current language is "+str(la)
                print(msg_para)
                document.add_paragraph(msg_para)
                
                text=v_info['text']
                self.Print_Info(text,"v_info['text']")
                x=translator.translate(text, dest=la)
                
                backup['text']=x.text
                s=str(backup)
                print(s)
                document.add_paragraph(s)
            
            if count>=lines:
                break
            
        li=headerList
        print(li)
        
        
        #For Header
        if li[0]==1:
            if li[2]==1:
                document=self.Auto_Add_Page_Number([0,li[1],0,""])
            if li[3]==1:
                document=self.Auto_Add_Page_Number([0,li[1],2,"/"])
            if li[4]==1:
                document=self.Auto_Add_Page_Number([0,li[1],1,""])
        
        li=footerList
        print(li)
        #For Footer
        if li[0]==1:
            if li[2]==1:
                document=self.Auto_Add_Page_Number([1,li[1],0,""])
            if li[3]==1:
                document=self.Auto_Add_Page_Number([1,li[1],2,"/"])
            if li[4]==1:       
                document=self.Auto_Add_Page_Number([1,li[1],1,""])
    
        
        self.document = document
        
        print("Almost complete the method named YT_Video_Transcript.")
        print("document_ext=")
        print(document_ext)
        for ext in document_ext:
            document.save(document_name+str(ext))

    
    def Auto_Add_Page_Number(self,options):
        
        WD_ALIGN_PARAGRAPH_list=[docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT
                                 ,docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
                                 , docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT]
        
        doc=self.document
        headerOrFooter=options[0]
        WD_ALIGN_PARAGRAPH_index=options[1]
        otherOptions=options[2:]
        if headerOrFooter==0:
            doc.sections[0].header.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH_list[WD_ALIGN_PARAGRAPH_index]
            self.add_page_number(run=doc.sections[0].header.paragraphs[0].add_run(),options=otherOptions)
        elif headerOrFooter==1:
            doc.sections[0].footer.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH_list[WD_ALIGN_PARAGRAPH_index]
            self.add_page_number(run=doc.sections[0].footer.paragraphs[0].add_run(),options=otherOptions)
        return doc
    
    def TryOpenUrl(self,base_channel_url,tar_channel_url,wantToOpen=1):
        if isinstance(wantToOpen, int):
            if wantToOpen==1:
                url=base_channel_url+tar_channel_url
                print("url=")
                print(url)
                webbrowser.open(url)
                return None
            return None
        else:
            raise TypeError

def main():
    
    print("sys.argv=")
    print(sys.argv)
    """
    for v in sys.argv:
        print(v)
        print(type(v))
        print("~~~~~~~~~~~~~`")
    """
    base_channel_url="https://www.youtube.com/channel/"

    #channel_url="UCO-n4ZDDXKPKK29c5eaytpA"
    exec_file_name=sys.argv[0]
    naming_method=str(sys.argv[1])
    tar_channel_url=str(sys.argv[2])
    document_name=str(sys.argv[3])
    document_ext=json.loads('"'+str(sys.argv[4])+'"')
    
    
    transcript_languages=json.loads('"'+str(sys.argv[5])+'"')
    #transcript_languages=['ja']
    
    tar_languanges=json.loads('"'+str(sys.argv[6])+'"')
    
    needHeader=str(sys.argv[7])
    headerFormatAlign=json.loads('"'+str(sys.argv[8])+'"')
    headerFormatCurrentPage=str(sys.argv[9])
    headerFormatSeperator=str(sys.argv[10])
    headerFormatTotalPage=str(sys.argv[11])
    
    needFooter=str(sys.argv[12])
    footerFormatAlign=json.loads('"'+str(sys.argv[13])+'"')
    footerFormatAlign=str(sys.argv[14])
    footerFormatSeperator=str(sys.argv[15])
    footerFormatTotalPage=str(sys.argv[16])
    
    fontSize=str(sys.argv[17])
    fontBody=str(sys.argv[18])
    fontTextColor=str(sys.argv[19])
    fontBackgroundColor=str(sys.argv[20])
    
    fontBold=str(sys.argv[21])
    fontUnderline=str(sys.argv[22])
    fontItalic=str(sys.argv[23])
    fontStrike=str(sys.argv[24])
    
    if tar_channel_url.startswith(base_channel_url):
        tar_channel_url=tar_channel_url.lstrip(base_channel_url) 
    
    headerList=[needHeader,headerFormatCurrentPage,headerFormatSeperator,headerFormatTotalPage,headerFormatAlign]
    footerList=[needFooter,footerFormatAlign,footerFormatSeperator,footerFormatTotalPage,footerFormatAlign]
    fontList1=[fontSize,fontBody,fontTextColor,fontBackgroundColor]
    fontList2=[fontBold,fontUnderline,fontItalic,fontStrike]
    
    wantToOpen=0
    
    #lines=1
    lines=10

    print(tar_channel_url)
    YT_video(transcript_languages=transcript_languages
             ,base_channel_url=base_channel_url
             ,tar_channel_url=tar_channel_url
             ,wantToOpen=wantToOpen
             ,document_name=document_name
             ,document_ext=document_ext
             ,naming_method=ORDER_class.ORDER.ASCENDING_FROM_1
             ,tar_languanges=tar_languanges
             ,lines=lines
             ,headerList=headerList
             ,footerList=footerList
             ,fontList1=fontList1
             ,fontList2=fontList2
             )
    

main()
