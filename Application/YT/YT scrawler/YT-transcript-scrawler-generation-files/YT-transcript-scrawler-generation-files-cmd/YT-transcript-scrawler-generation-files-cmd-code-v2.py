# import module
from __future__ import absolute_import

from lxml import etree

import re

import docx

from docx.compat import Unicode
from docx.oxml import OxmlElement
from docx.oxml.exceptions import InvalidXmlError
from docx.oxml.ns import NamespacePrefixedTag, nsmap, qn
from docx.shared import lazyproperty

from docx.oxml import ns

import sys

import json

import docx

from docx import shared

from docx import Document

from docx.shared import Inches

from docx import section

from docx.enum.section import WD_SECTION

from youtube_transcript_api import YouTubeTranscriptApi

from googletrans import Translator

import scrapetube

import webbrowser

import os

import ORDER_class
        
### ---------------------------------- ###
# Class: 
# A class to fetch some records of transcript given 
### ---------------------------------- ### 
class YT_video():
    
    ### ---------------------------------- ###
    # Function: 
    # Split a string to a list.
    # An utility method of HexString2DecNumber.
    ### ---------------------------------- ###
    def SplitDigit(self,s,many):
        li=[ s[idx:min(idx+many,len(s))] for idx in range(0,len(s),many)]
        return li
    
    ### ---------------------------------- ###
    # Function: 
    # Convert an object with string type with hex representation to a list of decimal number.
    # It is useful to convert a RGB color with string representation to a list type (i.e. for example [0,255,255]).
    ### ---------------------------------- ###
    def HexString2DecNumber(self,s,many):
        s=s.lstrip("#")
        numSymbolSign=s[0:2]
        s=s[2:]  
        li=self.SplitDigit(s=s,many=many)
        n_li=[ int(numSymbolSign+l_elem,0) for l_elem in li ]
        return n_li
    
    ### ---------------------------------- ###
    # Function: 
    # Convert an object which looks like a list with string type to a list.
    ### ---------------------------------- ###
    def StringToList(self,s:str):    
        x=s.strip('"')
        y=x.lstrip("[").rstrip("]").split(",")
        z=list()
        for yy in y:
            yyy=yy.strip("'").strip('"')
            z.append(yyy)
        return z
    
    ### ---------------------------------- ###
    # Function: 
    # Set an attribute to a default value.
    # An utility method of Set_All_DefaultValues.
    ### ---------------------------------- ###
    def Print_Info(self,o,msg):
        print(msg+"=")
        print(o)
        print("type("+msg+")=")
        print(type(o))
        
    ### ---------------------------------- ###
    # Function: 
    # Set the dictionary of class.
    # It is a special name.
    ### ---------------------------------- ###
    def __dict__(self):
        dic={
            "self.transcript_languages=":self.transcript_languages
            ,"self.base_channel_url=":self.base_channel_url
            ,"self.tar_channel_url=":self.tar_channel_url
            ,"self.wantToOpen=":self.wantToOpen
            ,"self.document_path=":self.document_path
            ,"self.document_name=":self.document_name
            ,"self.document_ext=":self.document_ext
            ,"self.naming_method=":self.naming_method
            ,"self.tar_languages=":self.tar_languages
            ,"self.lines=":self.lines
            ,"self.headerList=":self.headerList
            ,"self.footerList=":self.footerList
            ,"self.fontList1=":self.fontList1
            ,"self.fontList2=":self.fontList2
            ,"self.base_watch_video_url=":self.base_watch_video_url
            ,"self.base_channel_url=":self.base_channel_url
            ,"self.videos=":self.videos
            ,"self.video_url_list=":self.video_url_list
        }
        return dic
    
    ### ---------------------------------- ###
    # Function: 
    # Print all keys and values of the dictionary of the class.
    ### ---------------------------------- ###
    def Print_Dict_ElemByElem(self):
        print("The method of the class named Print_Dict_ElemByElem was called.")
        dic=self.__dict__()
        for k,v in dic.items():
            print("key:value,"+str(k)+":"+str(v))
            print("type(key):type(value),"+str(type(k))+":"+str(type(v)))
                    
    ### ---------------------------------- ###
    # Function: 
    # Constructor function.
    # __init__() is a special name.
    ### ---------------------------------- ###    
    def __init__(self
                 ,base_channel_url
                 ,tar_channel_url
                 ,wantToOpen
                 ,document_path: str
                 ,document_name: str
                 ,document_ext : list 
                 ,naming_method : ORDER_class.ORDER
                 ,transcript_languages
                 ,tar_languages : list
                 ,lines:int
                 ,headerList:list
                 ,footerList:list
                 ,fontList1:list
                 ,fontList2:list
                 ):
        base_watch_video_url='https://www.youtube.com/watch?v='
        self.base_watch_video_url=base_watch_video_url
        self.base_channel_url = base_channel_url
        
        self.tar_channel_url = tar_channel_url
        self.wantToOpen = wantToOpen
        self.document_path=document_path
        self.document_name = document_name
        self.document_ext = document_ext
        self.naming_method = naming_method
        
        self.transcript_languages=transcript_languages
        self.tar_languages = tar_languages
        self.lines=lines
        
        self.headerList=headerList
        self.footerList=footerList
        self.fontList1=fontList1
        self.fontList2=fontList2
        
        self.videos=list()
        self.video_url_list=list()
    
        self.TryOpenUrl(self.base_channel_url,self.tar_channel_url,self.wantToOpen)
        self.GetChannel()
        self.GetAllVideoUrl_List()
              
        self.Print_Dict_ElemByElem()       
        self.YT_handler()
        
    ### ---------------------------------- ###
    # Function: 
    # Set an attribute to a default value.
    # An utility method of Set_All_DefaultValues.
    ### ---------------------------------- ###
    def Set_DefaultValue(self,tar,defaultValue):
        if isinstance(tar,(list,tuple,set)):
            if len(tar)>=1:
                return tar
        
        if isinstance(tar,str):
            if tar!="" and len(tar)==0:
                return tar
        
        if isinstance(tar,(list,tuple)):
            tar=defaultValue
            return tar
        
        if isinstance(tar,(str,int,bool)):
            tar=defaultValue
            return tar
        raise TypeError
    
    ### ---------------------------------- ###
    # Function: 
    # Set All default values.
    # It is only useful when testing.
    # When I develop it, I have to use default value to check it is correct.
    ### ---------------------------------- ###
    def Set_All_DefaultValues(self):       
        self.transcript_languages = self.Set_DefaultValue(self.transcript_languages,['ja'])
        self.base_channel_url = self.Set_DefaultValue(self.base_channel_url,"https://www.youtube.com/channel/")
        self.tar_channel_url = self.Set_DefaultValue(self.tar_channel_url,"UCO-n4ZDDXKPKK29c5eaytpA")
        self.wantToOpen = self.Set_DefaultValue(self.wantToOpen,1)
        self.document_name = self.Set_DefaultValue(self.document_name,"Nanaten")
        self.document_ext = self.Set_DefaultValue(self.document_ext,[".docx"])
        self.naming_method = self.Set_DefaultValue(self.naming_method,ORDER_class.ORDER.ASCENDING_FROM_1)
        self.tar_languages = self.Set_DefaultValue(self.tar_languages,['en','zh-tw','ja'])
        
    ### ---------------------------------- ###
    # Function: 
    # Concatenate a list to string.
    ### ---------------------------------- ###
    def ConcatentateListToLetter(self,otherOptions):
        insertedText=""
        if len(otherOptions)<0:
            insertedText=""
            raise "Too less options in the variable otherOptions!!!"
            
        if len(otherOptions)==0:
            insertedText="/"
            return insertedText
        if len(otherOptions)==1:
            insertedText=otherOptions[0]
            return insertedText
        
        for s in otherOptions:
            insertedText=insertedText+str(s)
        return insertedText
    
    ### ---------------------------------- ###
    # Function: 
    # Create a new XML element.
    # An utility method of add_page_number.
    ### ---------------------------------- ###
    def create_element(self,name):
        return OxmlElement(name)

    ### ---------------------------------- ###
    # Function: 
    # Create anew XML attribute.
    # An utility method of add_page_number.
    ### ---------------------------------- ###
    def create_attribute(self,element, name, value):
        element.set(ns.qn(name), value)
    
    ### ---------------------------------- ###
    # Function: 
    # Add some texts in header or footer.
    ### ---------------------------------- ###
    def add_page_number(self,run,options):
        option0=options[0]
        
        # For other
        if option0!=0 and option0!=1:
            
            otherOptions=options[1:]
            insertedText=self.ConcatentateListToLetter(otherOptions)
            
            fldChar1 = self.create_element('w:fldChar')
            self.create_attribute(fldChar1, 'w:fldCharType', 'begin')
            
            fldChar2 = self.create_element('w:t')
            fldChar2.text=insertedText
            self.create_attribute(fldChar2, 'w:fldCharType', 'separate')

            fldChar3 = self.create_element('w:fldChar')
            self.create_attribute(fldChar2, 'w:fldCharType', 'end')
            
            run._r.append(fldChar1)
            run._r.append(fldChar2)
            run._r.append(fldChar3)
            
        # For current page and total page
        else:
            fldChar1 = self.create_element('w:fldChar')
            self.create_attribute(fldChar1, 'w:fldCharType', 'begin')
            
            instrText = self.create_element('w:instrText')
            self.create_attribute(instrText, 'xml:space', 'preserve')
        
        
            if option0==0:
                instrText.text = "PAGE"
            elif option0==1:
                instrText.text = "NUMPAGES"
    
            fldChar2 = self.create_element('w:fldChar')
            self.create_attribute(fldChar2, 'w:fldCharType', 'end')

            run._r.append(fldChar1)
            run._r.append(instrText)
            run._r.append(fldChar2)
    
    ### ---------------------------------- ###
    # Function: 
    # Get the YT channel ID with specified the YT channel link.
    ### ---------------------------------- ###
    def GetChannel(self):
        videos = scrapetube.get_channel(self.tar_channel_url)
        setattr(self, 'videos', videos)
    
    ### ---------------------------------- ###
    # Function: 
    # Get links for all videos.
    ### ---------------------------------- ###
    def GetAllVideoUrl_List(self):
        #self.video_url_list=list()
        v_list=list()
        for video in self.videos:
            v_url=video['videoId']
            video_url=str(self.base_watch_video_url)+v_url
            print("video_url=")
            print(video_url)
            v_list.append(video_url)
        setattr(self, 'video_url_list', v_list)
        
    ### ---------------------------------- ###
    # Function: 
    # To check the s with string type is valid for color with hex representation.
    # A valid string should be such as #0x112233 or #0X112233
    ### ---------------------------------- ###  
    def IsValidColorHexGRB(self,s:str,errorMsg:str):
        
        if (not s.startswith("#0x")) and (not s.startswith("#0X")):
            raise ValueError("Error!!! A valid color with hex representation must start with 0#x or #0X !!!")
            
        if len(s)>9:
            raise ValueError("Error!!! The length of "+errorMsg+" is too long!!!")
            
        if len(s)<9:
            raise ValueError("Error!!! The length of "+errorMsg+" is too short!!!")
            
        for idx in range(3,len(s),1):
            c=s[idx]
            
            #Convert the letter c as a int object through Unicode.
            unic=ord(c)
            
            # Check the unic is a digit by Unicode code.
            if not((48<=unic and unic<=57 ) or (97<=unic and unic<=101) or (65<=unic and unic<=69)):
                raise ValueError("Error!!! There exists an invalid letter of a valid color with hex representation !!!")
        
     ### ---------------------------------- ###
     # Function: 
     # To check the s with string type is valid for color with hex representation.
     # An utility method to YT_handler.
     ### ---------------------------------- ### 
    def YT_handler(self):   
        ## ---------------------------------- ## 
        # DON'T USE this statement.
        #for video_url_elem in self.videos: 
        # DON'T directly use the attribute of the class since it is generation expression.
        # Convert it to the desired type before used.
        
        # DON'T USE this statement.
        # for video_url_elem in list(self.videos):
        # Since list(self.videos) is a generation expression and it returns an empty list.
        ## ---------------------------------- ## 
        
        ## Such as this statement
        cnt=0
        for video_url_elem in self.__dict__()['self.video_url_list=']:
            cnt+=1
            
            """
            Error usage
            
            self.YT_Video_Transcript(video_url=video_url_elem,
                                         document_name=self.document_name+str(cnt),
                                         document_ext=self.document_ext,
                                         transcript_language=self.transcript_languages,
                                         tar_languages=self.tar_languages,
                                         lines=self.lines
                                         ,headerList=self.headerList
                                         ,footerList=self.footerList
                                         ,fontList1=self.fontList1
                                         ,fontList2=self.fontList2                                                                                                             
                                         )
            """
            
            self.YT_Video_Transcript(video_url=video_url_elem,
                                         document_path=self.__dict__()['self.document_path='],
                                         document_name=self.__dict__()['self.document_name=']+str(cnt),
                                         
                                         document_ext=self.__dict__()['self.document_ext='],
                                         transcript_language=self.__dict__()['self.transcript_languages='],
                                         tar_languages= self.__dict__()['self.tar_languages='],
                                         lines= self.__dict__()['self.lines=']
                                         ,headerList= self.__dict__()['self.headerList=']
                                         ,footerList= self.__dict__()['self.footerList=']
                                         ,fontList1=self.__dict__()['self.fontList1=']
                                         ,fontList2=self.__dict__()['self.fontList2=']
                                         )
            
    ### ---------------------------------- ###
    # Function: 
    # To check the s with string type is valid for color with hex representation.
    # An utility method to YT_handler.
    ### ---------------------------------- ### 
    def YT_Video_Transcript(self
                            ,video_url:str
                            ,document_path:str
                            ,document_name:str
                            ,document_ext:str
                            ,transcript_language 
                            ,tar_languages
                            ,lines:int
                            ,headerList
                            ,footerList
                            ,fontList1
                            ,fontList2
                            ):
        
        # Create a new empty Translator Object.
        translator = Translator()
        
        # Create a new empty Document Object.
        document = Document()
        
        # Unpack the list
        fontSize=fontList1[0]
        fontBody=fontList1[1]
        fontTextColor=fontList1[2]
        fontTextBackgroundColor=fontList1[3]
        
        fontTextColor=fontTextColor.strip('"').strip("'")
        
        # Check validity of the fontTextColor.
        self.IsValidColorHexGRB(fontTextColor, "fontTextColor")
        fontTextColor=self.HexString2DecNumber(fontTextColor, 2) 
        fontTextColor_r=fontTextColor[0]
        fontTextColor_g=fontTextColor[1]
        fontTextColor_b=fontTextColor[2]
        
        # Check validity of the fontTextBackgroundColor.
        fontTextBackgroundColor=fontTextBackgroundColor.strip('"').strip("'")
        self.IsValidColorHexGRB(fontTextBackgroundColor, "fontTextBackgroundColor")   
        fontTextBackgroundColor=self.HexString2DecNumber(fontTextBackgroundColor, 2)
        
        
        # Convert it to a list which, for each elements, it is a bool type.
        fontList2=[True if fontList2[x]=='1' else False for x in range(0,len(fontList2),1)]
        
        headerAlignList=headerList[-1]
        headerAlignList=self.StringToList(headerAlignList)
        headerList=[True if headerList[x]=='1' else False for x in range(0,len(headerList)-1,1)]
        
        footerAlignList=footerList[-1]
        footerAlignList=self.StringToList(footerAlignList)
        footerList=[True if footerList[x]=='1' else False for x in range(0,len(footerList)-1,1)]
        
        # Convert an object which looks like a list with string type to a list.
        transcript_language=self.StringToList(transcript_language)
        
        ## ----------------------------------- ##
        # Check the validity of the variable transcript_language.
        # if transcript_language is NOT a string and its lens is NOT equal to 1.
        # Then raise ValueError.
        ## ----------------------------------- ##
        if not isinstance(transcript_language,(str)) and len(transcript_language)!=1:
            raise ValueError("Error!!! Invalid value for the variable transcript_language!!!")
            
        
        ## ----------------------------------- ##
        # To fecth all records of transcripts of specified YT video link. 
        # (NOTE that video)
        ## ----------------------------------- ##
        video_id=video_url.split('=')[1]
        transcript_list = YouTubeTranscriptApi.list_transcripts(video_id)
        transcript=transcript_list.find_generated_transcript(transcript_language)
        video_info=transcript.fetch()
        
        start_index=7
        start_index=start_index+1
        
        msg="The transcript of YT, link="+str(video_url)
        
        # Add a new paragraph.
        p=document.add_paragraph()
        
        # Add a new run with specified text named msg.
        run=p.add_run(msg)
        
        # Determine to set the font of run as bold text or NOT.
        run.font.bold=fontList2[0]
        
        # Determine to set the font of run as underline text or NOT.
        run.font.underline=fontList2[1]
        
        # Determine to set the font of run as italic text or NOT.
        run.font.italic=fontList2[2]
        
        # Determine to set the font of run as strike text or NOT.
        run.font.strike=fontList2[3]

        # Set the size of font.
        run.font.size = docx.shared.Pt(int(fontSize))
        
        # Set the font family of text.
        fontBody=fontBody.strip('"').strip("'")
        run.font.name = fontBody

        # Set the color of text.
        run.font.color.rgb=docx.shared.RGBColor(fontTextColor_r,fontTextColor_g,fontTextColor_b)
        
        ##---------------##
        # To set the background color of text with any valid specified arbitrary values.
        # We have to insert an XML tag inside the block to replace the default value.
        # It requires the concept of XML.
        
        # We can NOT use the following statements.
        # Otherwise, error occurs at runtime since there are NO attribute rgb in the run.font.highlight_color
        # run.font.highlight_color.rgb=docx.shared.RGBColor(fontTextBackgroundColor[0],fontTextBackgroundColor[1],fontTextBackgroundColor[2]) 
        ##---------------##
        
        # Get the XML tag
        tag = run._r

        # Create XML element
        shd = OxmlElement('w:shd')

        # Add attributes to the element
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), fontTextBackgroundColor)
        
        tag.rPr.append(shd)
          
        # Append a page break.
        document.add_page_break()

        # Convert an object which looks like a list but with string type to a list with list type
        tar_languages=self.StringToList(tar_languages)
          
        #self.Print_Info(fontList2,"fontList2")        
        count=0
        for v_info in video_info:
      
            count+=1
            for la in tar_languages:
                
                # Assign v_info to other variable to keep the original data of v_info. 
                backup=v_info
                
                msg_para="The current language is "+str(la)
                
                p=document.add_paragraph()
                run=p.add_run(msg_para)
                run.font.bold=fontList2[0]
                run.font.underline=fontList2[1]
                run.font.italic=fontList2[2]
                run.font.strike=fontList2[3]
                run.font.size = docx.shared.Pt(int(fontSize))                
                fontBody=fontBody.strip('"').strip("'")
                run.font.color.rgb=docx.shared.RGBColor(fontTextColor_r,fontTextColor_g,fontTextColor_b)
                run.font.name = fontBody
                tag = run._r
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), fontTextBackgroundColor)
                tag.rPr.append(shd)
                
                text=v_info['text']
               
                # Translate the text into the specified language named la.
                x=translator.translate(text, dest=la)
                
                # Assign the text of translated language.
                backup['text']=x.text
                
                s=str(backup)
                p=document.add_paragraph()
                run=p.add_run(s)
                run.font.bold=fontList2[0]
                run.font.underline=fontList2[1]
                run.font.italic=fontList2[2]
                run.font.strike=fontList2[3]
                run.font.size = docx.shared.Pt(int(fontSize))
                fontBody=fontBody.strip('"').strip("'")
                run.font.name = fontBody
                run.font.color.rgb=docx.shared.RGBColor(fontTextColor_r,fontTextColor_g,fontTextColor_b) 
                tag = run._r
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), fontTextBackgroundColor)     
                tag.rPr.append(shd)
            
            # When True, it indicates the number of records which are fetched achieve the maximum.
            if count>=lines:
                break
            
        li=headerList
        #For Header
        if li[0]==1:
            for k in ['0','1','2']:
                if k in headerAlignList:
                    n=int(k)
                    self.Print_Info(n, "n")
                    if li[1]==True:
                        document=self.Auto_Add_Page_Number(document=document,options=[0,n,0,""])
                    if li[2]==True:
                        document=self.Auto_Add_Page_Number(document=document,options=[0,n,2,"/"])
                    if li[3]==True:
                        document=self.Auto_Add_Page_Number(document=document,options=[0,n,1,""])
        
        li=footerList
        #For Footer
        if li[0]==1:
            for k in ['0','1','2']:
                if k in footerAlignList:
                    n=int(k)
                    self.Print_Info(n, "n")
                    if li[1]==True:
                        document=self.Auto_Add_Page_Number(document=document,options=[1,n,0,""])
                    if li[2]==True:
                        document=self.Auto_Add_Page_Number(document=document,options=[1,n,2,"/"])
                    if li[3]==True:       
                        document=self.Auto_Add_Page_Number(document=document,options=[1,n,1,""])
        
        document_ext=self.StringToList(document_ext)
        for ext in document_ext:
            
            # Join path, filename, and document extension to a fullname.
            d_name=os.path.join(document_path,document_name)+str(ext)
            
            # Save the document which name is d_name.
            document.save(d_name)
            
    ### ---------------------------------- ###
    # Function: 
    # Auto add page number for headers and footers by specified options.
    ### ---------------------------------- ### 
    def Auto_Add_Page_Number(self,document,options):
        
        WD_ALIGN_PARAGRAPH_list=[docx.enum.text.WD_ALIGN_PARAGRAPH.LEFT
                                 ,docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
                                 , docx.enum.text.WD_ALIGN_PARAGRAPH.RIGHT]
        doc=document
        headerOrFooter=options[0]
        WD_ALIGN_PARAGRAPH_index=options[1]
        otherOptions=options[2:]
        if headerOrFooter==0:
            doc.sections[0].header.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH_list[WD_ALIGN_PARAGRAPH_index]
            self.add_page_number(run=doc.sections[0].header.paragraphs[0].add_run(),options=otherOptions)
        elif headerOrFooter==1:
            doc.sections[0].footer.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH_list[WD_ALIGN_PARAGRAPH_index]
            self.add_page_number(run=doc.sections[0].footer.paragraphs[0].add_run(),options=otherOptions)
        return doc
    
    ### ---------------------------------- ###
    # Function: 
    # Try open the url with specified YT channel link.
    ### ---------------------------------- ### 
    def TryOpenUrl(self,base_channel_url,tar_channel_url,wantToOpen=1):
        if isinstance(wantToOpen, int):
            if wantToOpen==1:
                url=base_channel_url+tar_channel_url
                print("url=")
                print(url)
                webbrowser.open(url)
                return None
            return None
        else:
            raise TypeError

### ---------------------------------- ###
# Function: 
# Driver main function.    
### ---------------------------------- ### 
def main():
    
    ### ---------------------------------- ###
    # Function: 
    # To print the info of o.
    ### ---------------------------------- ### 
    def Print_Info(o,msg):
        print(msg+"=")
        print(o)
        print("type("+msg+")=")
        print(type(o))
        print(tar_channel_url)
        
    print("sys.argv=")
    print(sys.argv)
    base_channel_url="https://www.youtube.com/channel/"

    #channel_url="UCO-n4ZDDXKPKK29c5eaytpA"
    
    exec_file_name=sys.argv[0]
    naming_method=str(sys.argv[1])
    tar_channel_url=str(sys.argv[2])
    
    document_path=str(sys.argv[3]).strip('"').strip("'")
    document_name=str(sys.argv[4])
    document_ext=json.loads('"'+str(sys.argv[5])+'"')
    
    
    transcript_languages=json.loads('"'+str(sys.argv[6])+'"')
    
    transcript_languages='"'+transcript_languages+'"'   
    tar_languages=json.loads('"'+str(sys.argv[7])+'"')
    
    needHeader=str(sys.argv[8])
    headerFormatAlign=json.loads('"'+str(sys.argv[9])+'"')
    headerFormatCurrentPage=str(sys.argv[10])
    headerFormatSeperator=str(sys.argv[11])
    headerFormatTotalPage=str(sys.argv[12])
    
    needFooter=str(sys.argv[13])
    footerFormatAlign=json.loads('"'+str(sys.argv[14])+'"')
    footerFormatCurrentPage=str(sys.argv[15])
    footerFormatSeperator=str(sys.argv[16])
    footerFormatTotalPage=str(sys.argv[17])
    
    fontSize=str(sys.argv[18])
    fontBody=str(sys.argv[19])
    fontTextColor=str(sys.argv[20])
    fontTextBackgroundColor=str(sys.argv[21])
    
    fontBold=str(sys.argv[22])
    fontUnderline=str(sys.argv[23])
    fontItalic=str(sys.argv[24])
    fontStrike=str(sys.argv[25])
    
    if tar_channel_url.startswith(base_channel_url):
        tar_channel_url=tar_channel_url.lstrip(base_channel_url) 
    
    headerList=[needHeader,headerFormatCurrentPage,headerFormatSeperator,headerFormatTotalPage,headerFormatAlign]
    footerList=[needFooter,footerFormatCurrentPage,footerFormatSeperator,footerFormatTotalPage,footerFormatAlign]
    fontList1=[fontSize,fontBody,fontTextColor,fontTextBackgroundColor]
    fontList2=[fontBold,fontUnderline,fontItalic,fontStrike]
    
    wantToOpen=0

    #lines=1
    lines=10

    YT_video(
             base_channel_url=base_channel_url
             ,tar_channel_url=tar_channel_url
             ,wantToOpen=wantToOpen
             ,document_path=document_path
             ,document_name=document_name
             ,document_ext=document_ext
             ,naming_method=ORDER_class.ORDER.ASCENDING_FROM_1
             ,transcript_languages=transcript_languages
             ,tar_languages=tar_languages
             ,lines=lines
             ,headerList=headerList
             ,footerList=footerList
             ,fontList1=fontList1
             ,fontList2=fontList2
             )
    
# Call the driver function.
main()
